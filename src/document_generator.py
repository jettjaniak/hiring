"""
Document generation utilities for filling DOCX and XLSX templates
"""
import os
from pathlib import Path
from typing import Dict, List
from docx import Document
from openpyxl import load_workbook
import re
from io import BytesIO


def get_template_path(template_filename: str) -> Path:
    """Get the absolute path to a template file"""
    base_dir = Path(__file__).parent.parent
    return base_dir / "document_templates" / template_filename


def fill_docx_template(template_filename: str, replacements: Dict[str, str]) -> BytesIO:
    """
    Fill a DOCX template with the provided replacements.

    Args:
        template_filename: Name of the template file (e.g., "offer_letter_template.docx")
        replacements: Dictionary of placeholder -> value mappings
                     (e.g., {"{{CANDIDATE_NAME}}": "John Doe"})

    Returns:
        BytesIO object containing the filled document
    """
    template_path = get_template_path(template_filename)

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                # Replace in runs to preserve formatting
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def fill_xlsx_template(template_filename: str, replacements: Dict[str, str]) -> BytesIO:
    """
    Fill an XLSX template with the provided replacements.

    Args:
        template_filename: Name of the template file (e.g., "background_check_template.xlsx")
        replacements: Dictionary of placeholder -> value mappings
                     (e.g., {"{{CANDIDATE_NAME}}": "John Doe"})

    Returns:
        BytesIO object containing the filled spreadsheet
    """
    template_path = get_template_path(template_filename)

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    wb = load_workbook(template_path)
    ws = wb.active

    # Replace in all cells
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                for key, value in replacements.items():
                    if key in cell.value:
                        cell.value = cell.value.replace(key, value)

    # Save to BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def extract_placeholders_from_docx(template_filename: str) -> List[str]:
    """
    Extract all placeholders ({{PLACEHOLDER}}) from a DOCX template.

    Args:
        template_filename: Name of the template file

    Returns:
        List of unique placeholder names found in the template
    """
    template_path = get_template_path(template_filename)

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    placeholders = set()

    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        matches = re.findall(r'\{\{([A-Z_]+)\}\}', paragraph.text)
        placeholders.update(matches)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(r'\{\{([A-Z_]+)\}\}', cell.text)
                placeholders.update(matches)

    return sorted(list(placeholders))


def extract_placeholders_from_xlsx(template_filename: str) -> List[str]:
    """
    Extract all placeholders ({{PLACEHOLDER}}) from an XLSX template.

    Args:
        template_filename: Name of the template file

    Returns:
        List of unique placeholder names found in the template
    """
    template_path = get_template_path(template_filename)

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    wb = load_workbook(template_path)
    ws = wb.active
    placeholders = set()

    # Extract from all cells
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                matches = re.findall(r'\{\{([A-Z_]+)\}\}', cell.value)
                placeholders.update(matches)

    return sorted(list(placeholders))
